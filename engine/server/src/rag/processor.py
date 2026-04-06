"""
Document processor for RAG ingestion.

Handles text extraction from various file types, chunking, and embedding generation.
Ported from backend/src/rag/processor/ for local runtime use.
"""

import asyncio
import logging
import re
from pathlib import Path
from uuid import uuid4

from src.embedding.resolver import get_knowledge_embedding_provider
from src.infra.config import get_settings

from .chunker import Chunk  # noqa: F401
from .chunker import RecursiveChunker as TextChunker  # noqa: F401

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".markdown"}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB


# ─── Text Extraction ──────────────────────────────────────────────────────────


async def extract_text(file_content: bytes, filename: str) -> str:
    """Extract text from file content based on extension."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return await asyncio.wait_for(asyncio.to_thread(extract_pdf, file_content), timeout=120.0)
    elif ext == ".docx":
        return await asyncio.wait_for(asyncio.to_thread(extract_docx, file_content), timeout=120.0)
    elif ext in (".txt", ".text"):
        return file_content.decode("utf-8", errors="replace")
    elif ext in (".md", ".markdown"):
        return strip_markdown(file_content.decode("utf-8", errors="replace"))
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def extract_pdf(content: bytes) -> str:
    """Extract text from PDF bytes.

    Tries pypdf first, falls back to pdfplumber if unavailable.
    Raises ImportError with a clear message if neither is installed.
    """
    import io

    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from pdfplumber import open as open_pdf
        except ImportError:
            raise ImportError(
                "PDF extraction requires either 'pypdf' or 'pdfplumber'. "
                "Install one with: pip install pypdf"
            ) from None

        pdf = open_pdf(io.BytesIO(content))
        parts = []
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text:
                parts.append(f"[Page {i}]\n{text}")
        return "\n\n".join(parts)

    reader = PdfReader(io.BytesIO(content))
    parts = []
    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text:
            parts.append(f"[Page {i}]\n{text}")
    return "\n\n".join(parts)


def extract_docx(content: bytes) -> str:
    """Extract text from DOCX bytes."""
    import io

    from docx import Document

    doc = Document(io.BytesIO(content))
    parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                parts.append(" | ".join(row_text))

    return "\n\n".join(parts)


def strip_markdown(text: str) -> str:
    """Strip markdown formatting from text."""
    text = re.sub(r"#{1,6}\s+", "", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    text = re.sub(r"```[\s\S]*?```", "", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


# ─── Processing Pipeline ──────────────────────────────────────────────────────


async def process_document(
    document_id: str,
    collection_id: str,
    file_content: bytes,
    filename: str,
    db_session,
    chunk_size: int = 500,
    chunk_overlap: int = 50,
) -> int:
    """Full document processing pipeline: extract → chunk → embed → store.

    Returns the number of chunks created.
    """
    from sqlalchemy import select

    from .models import RAGChunk, RAGCollection, RAGDocument

    get_settings()

    # 1. Extract text
    text = await extract_text(file_content, filename)
    if not text.strip():
        raise ValueError("No text content could be extracted from the document")

    # 2. Chunk text — use strategy from collection metadata if available
    from .chunker import ChunkerFactory

    # Determine chunking strategy from collection metadata
    coll_meta_result = await db_session.execute(
        select(RAGCollection.meta).where(RAGCollection.id == collection_id)
    )
    coll_meta = coll_meta_result.scalar_one_or_none() or {}
    if isinstance(coll_meta, dict):
        chunk_strategy = coll_meta.get("chunk_strategy", "token_aware")
    else:
        chunk_strategy = "token_aware"

    # Pre-create embedding provider for semantic chunking (and later embedding)
    embedding_provider = get_knowledge_embedding_provider()

    try:
        chunker = ChunkerFactory.get_chunker(
            chunk_strategy,
            embedding_provider=embedding_provider if chunk_strategy == "semantic" else None,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            chunk_size_tokens=chunk_size // 2,  # rough char-to-token ratio
            overlap_tokens=chunk_overlap // 2,
        )
    except ValueError:
        # Fallback to recursive chunker
        chunker = TextChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    chunks = chunker.split(text)

    if not chunks:
        raise ValueError("Document produced no chunks after splitting")

    # 3. Generate embeddings (batch, no Redis cache for document chunks)
    # embedding_provider already created above (reused for semantic chunking)
    texts = [c.content for c in chunks]

    _BATCH_SIZE = 100
    embeddings: list[list[float]] = []
    for i in range(0, len(texts), _BATCH_SIZE):
        batch = texts[i : i + _BATCH_SIZE]
        batch_embeddings = await embedding_provider.embed_texts(batch)
        embeddings.extend(batch_embeddings)

    # 4. Store chunk metadata in PG (no embedding column)
    chunk_ids: list[str] = []
    for _i, chunk in enumerate(chunks):
        chunk_id = str(uuid4())
        chunk_ids.append(chunk_id)
        db_chunk = RAGChunk(
            id=chunk_id,
            document_id=document_id,
            collection_id=collection_id,
            content=chunk.content,
            chunk_index=chunk.position,
            meta=chunk.metadata,
        )
        db_session.add(db_chunk)

    # 5. Upsert vectors + payloads to Qdrant
    from .vector_store import ChunkData, QdrantRAGVectorStore

    # Fetch parent collection scope info for Qdrant payloads
    collection_obj = await db_session.execute(
        select(RAGCollection).where(RAGCollection.id == collection_id)
    )
    coll = collection_obj.scalar_one_or_none()
    scope = coll.scope.value if coll else "global"
    group_slugs = list(coll.allowed_groups) if coll else []
    agent_id = coll.owner_user_id if coll and coll.scope.value == "agent" else None

    qdrant_chunks = [
        ChunkData(
            id=chunk_ids[i],
            content=chunks[i].content,
            embedding=embeddings[i],
            scope=scope,
            group_slugs=group_slugs,
            agent_id=agent_id,
            user_id=None,
            document_id=document_id,
            collection_id=collection_id,
            chunk_index=chunks[i].position,
            parent_chunk_id=getattr(chunks[i], "parent_id", None),
            chunk_level=getattr(chunks[i], "chunk_level", 0),
            metadata=getattr(chunks[i], "metadata", {}),
        )
        for i in range(len(chunks))
    ]

    vector_store = QdrantRAGVectorStore()
    await vector_store.upsert_chunks(qdrant_chunks)

    # 6. Update document chunk count
    from sqlalchemy import select, update

    await db_session.execute(
        update(RAGDocument).where(RAGDocument.id == document_id).values(chunk_count=len(chunks))
    )

    # 6. Update collection chunk/document counts
    #    Flush first so the count queries include the newly added chunks.
    await db_session.flush()

    from sqlalchemy import func

    total_chunks = (
        await db_session.execute(
            select(func.count(RAGChunk.id)).where(RAGChunk.collection_id == collection_id)
        )
    ).scalar() or 0
    total_docs = (
        await db_session.execute(
            select(func.count(RAGDocument.id)).where(RAGDocument.collection_id == collection_id)
        )
    ).scalar() or 0

    await db_session.execute(
        update(RAGCollection)
        .where(RAGCollection.id == collection_id)
        .values(chunk_count=total_chunks, document_count=total_docs)
    )

    await db_session.flush()

    return len(chunks)
